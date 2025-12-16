from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional
import shutil
import os
import uuid
import datetime

from database import get_db, Transcript, QAEntry, ProcessingStatus
from services import UPLOAD_DIR

router = APIRouter()

os.makedirs(UPLOAD_DIR, exist_ok=True)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RECORDED_DIR = os.path.join(BASE_DIR, "../data/recorded")
os.makedirs(RECORDED_DIR, exist_ok=True)

@router.get("/health")
def health_check():
    return {"status": "ok"}

@router.post("/upload")
async def upload_audio(
    file: UploadFile = File(...), 
    db: Session = Depends(get_db)
):
    existing = db.query(Transcript).filter(Transcript.filename == file.filename).first()
    if existing:
        raise HTTPException(status_code=409, detail=f"File '{file.filename}' sudah ada di sistem.")

    file_id = str(uuid.uuid4())
    filename = file.filename
    ext = os.path.splitext(filename)[1]
    
    safe_name = "".join([c for c in filename if c.isalpha() or c.isdigit() or c in (' ', '.', '_', '-')]).rstrip()
    if not safe_name:
        safe_name = f"upload_{file_id}{ext}"
        
    save_path = os.path.abspath(os.path.join(UPLOAD_DIR, safe_name))
    
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    new_transcript = Transcript(
        id=file_id,
        filename=filename,
        file_path=save_path,
        status=ProcessingStatus.PENDING,
        current_step="Uploaded. Waiting for approval..."
    )
    db.add(new_transcript)
    db.commit()
    db.refresh(new_transcript)
    
    return {"message": "File uploaded", "id": file_id}

@router.post("/record")
async def save_recording(
    file: UploadFile = File(...),
    filename: str = Form(...),
    transcribe: bool = Form(True),
    db: Session = Depends(get_db)
):
    file_id = str(uuid.uuid4())
    safe_filename = "".join([c for c in filename if c.isalpha() or c.isdigit() or c in (' ', '.', '_')]).rstrip()
    if not safe_filename.endswith(('.wav', '.webm', '.mp3')):
         safe_filename += ".wav"
         
    save_path = os.path.join(RECORDED_DIR, f"{safe_filename}")
    
    if os.path.exists(save_path):
        ts = datetime.datetime.now().strftime("%H%M%S")
        save_path = os.path.join(RECORDED_DIR, f"{ts}_{safe_filename}")

    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    if transcribe:
        new_transcript = Transcript(
            id=file_id,
            filename=os.path.basename(save_path),
            file_path=save_path,
            status=ProcessingStatus.QUEUED,
            current_step="Queued from recording..."
        )
        db.add(new_transcript)
        db.commit()
        return {"message": "Recording saved and queued", "id": file_id}
    else:
        return {"message": "Recording saved", "path": save_path}

@router.get("/transcripts")
def list_transcripts(db: Session = Depends(get_db)):
    return db.query(Transcript).order_by(Transcript.upload_date.desc()).all()

@router.get("/transcripts/{t_id}")
def get_transcript(t_id: str, db: Session = Depends(get_db)):
    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t:
        raise HTTPException(404, "Transcript not found")
    return t

@router.post("/transcripts/{t_id}/retry")
def retry_transcript(t_id: str, db: Session = Depends(get_db)):
    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t:
        raise HTTPException(404, "Transcript not found")
    
    t.status = ProcessingStatus.QUEUED
    t.progress = 0
    t.current_step = "Retrying..."
    db.commit()
    t.current_step = "Retrying..."
    db.commit()
    return {"message": "Retrying"}

@router.post("/transcripts/{t_id}/start")
def start_transcript(t_id: str, db: Session = Depends(get_db)):
    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t:
        raise HTTPException(404, "Transcript not found")
    
    t.status = ProcessingStatus.QUEUED
    t.current_step = "Queued for processing..."
    db.commit()
    return {"message": "Transcription Started"}

@router.delete("/transcripts/{t_id}")
def delete_transcript(t_id: str, db: Session = Depends(get_db)):
    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t:
        raise HTTPException(404, "Transcript not found")
        
    if t.file_path and os.path.exists(t.file_path):
        try:
            os.remove(t.file_path)
        except Exception as e:
            print(f"Error deleting file {t.file_path}: {e}")

    db.query(QAEntry).filter(QAEntry.transcript_id == t_id).delete()
    
    db.delete(t)
    db.commit()
    
    return {"message": "Deleted successfully"}

@router.delete("/qa/{qa_id}")
def delete_qa_entry(qa_id: str, db: Session = Depends(get_db)):

    entry = db.query(QAEntry).filter(QAEntry.id == qa_id).first()
    if not entry:
        raise HTTPException(404, "Q&A Entry not found")
    
    db.delete(entry)
    db.commit()
    return {"message": "Deleted successfully"}

from docx import Document
from fastapi.responses import FileResponse
import tempfile

@router.get("/transcripts/{t_id}/download_qa")
def download_qa_docx(t_id: str, db: Session = Depends(get_db)):

    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t:
        raise HTTPException(404, "Transcript not found")
        
    qa_list = db.query(QAEntry).filter(QAEntry.transcript_id == t_id).all()
    
    if not qa_list:
        raise HTTPException(400, "No Q&A data to download")

    doc = Document()
    doc.add_heading(f'Laporan Analisis: {t.filename}', 0)
    
    doc.add_paragraph(f"ID Transkrip: {t.id}")
    doc.add_paragraph(f"Tanggal: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("-" * 50)
    
    doc.add_heading('Transkrip Lengkap', level=1)
    if t.raw_text:
        doc.add_paragraph(t.raw_text)
    else:
        doc.add_paragraph("[Transkrip belum tersedia/kosong]")
    
    doc.add_page_break()
    
    doc.add_heading('Data Pasien (Ekstraksi Otomatis)', level=1)
    
    for idx, qa in enumerate(qa_list, 1):
        label = qa.question.split(":", 1)[0] if ":" in qa.question else qa.question
        
        doc.add_heading(f'{label}', level=2)

        
        p_a = doc.add_paragraph()
        p_a.add_run("A: ").bold = True
        p_a.add_run(qa.answer or "Belum dijawab")
        
        if qa.bleu_score is not None:
             p_s = doc.add_paragraph()
             p_s.add_run(f"Confidence Score: {qa.bleu_score:.4f}").italic = True
             p_s.style = "Caption"
        
        if qa.context_used:
             p_ctx = doc.add_paragraph()
             p_ctx.add_run("Konteks: ").bold = True
             p_ctx.add_run(qa.context_used[:500] + "...")
             
        doc.add_paragraph("_" * 20)

    temp_filename = f"Laporan_QA_{t.filename}.docx"
    safe_name = "".join([c for c in temp_filename if c.isalpha() or c.isdigit() or c in (' ', '.', '_', '-')]).rstrip()
    
    temp_path = os.path.join(tempfile.gettempdir(), safe_name)
    doc.save(temp_path)
    
    return FileResponse(temp_path, filename=safe_name, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@router.get("/transcripts/{t_id}/download_text")
def download_transcript_text(t_id: str, db: Session = Depends(get_db)):
    t = db.query(Transcript).filter(Transcript.id == t_id).first()
    if not t or not t.raw_text:
        raise HTTPException(404, "Transcript text not available")
    
    temp_filename = f"{t.filename}.txt"
    safe_name = "".join([c for c in temp_filename if c.isalpha() or c.isdigit() or c in (' ', '.', '_', '-')]).rstrip()
    temp_path = os.path.join(tempfile.gettempdir(), safe_name)
    
    with open(temp_path, "w", encoding="utf-8") as f:
        f.write(t.raw_text)
        
    return FileResponse(temp_path, filename=safe_name, media_type='text/plain')



@router.get("/qa/{t_id}")
def get_qa_results(t_id: str, db: Session = Depends(get_db)):
    return db.query(QAEntry).filter(QAEntry.transcript_id == t_id).all()
